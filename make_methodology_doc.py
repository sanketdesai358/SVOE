"""Generate SVOE Methodology Word document using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from pathlib import Path

OUTPUT = Path("SVOE_Methodology_v2.docx")

# ── colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x1A, 0x5E)
BLUE   = RGBColor(0x21, 0x96, 0xF3)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, style in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style.get("val", "single"))
        el.set(qn("w:sz"), str(style.get("sz", 4)))
        el.set(qn("w:color"), style.get("color", "CCCCCC"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color="2196F3", thickness=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY if level == 1 else DARK
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.name = "Arial"
    if level == 1:
        add_horizontal_rule(doc, color="2196F3", thickness=8)
    return p


def body(doc, text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Arial"
    return p


def formula_box(doc, formula: str, caption: str = ""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "EFF8FF")
    set_cell_border(cell,
        top={"val": "single", "sz": 6, "color": "2196F3"},
        bottom={"val": "single", "sz": 6, "color": "2196F3"},
        left={"val": "single", "sz": 6, "color": "2196F3"},
        right={"val": "single", "sz": 6, "color": "2196F3"},
    )
    cell._tc.get_or_add_tcPr()
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(formula)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Courier New"
    r.font.color.rgb = NAVY
    if caption:
        doc.add_paragraph()
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GREY
        cr.font.name = "Arial"
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(10)


def data_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "1A1A5E")
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(4)
        r = cp.add_run(hdr)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = "Arial"

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(3)
            cp.paragraph_format.space_after = Pt(3)
            r = cp.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "Arial"

    # Column widths
    if col_widths:
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── page margins (1 inch all sides, US Letter) ────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── default font ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER / TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(36)
title_p.paragraph_format.space_after = Pt(6)
tr = title_p.add_run("Shot Value Over Expected (SVOE)")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = NAVY
tr.font.name = "Arial"

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
sr = sub_p.add_run("Methodology & Technical Documentation")
sr.font.size = Pt(14)
sr.font.color.rgb = BLUE
sr.font.name = "Arial"

# Rule under title
add_horizontal_rule(doc, color="2196F3", thickness=16)

# Meta line
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(2)
mr = meta_p.add_run("NBA Analytics  |  2022-23 through 2025-26 Seasons  |  XGBoost Model")
mr.font.size = Pt(10)
mr.font.color.rgb = GREY
mr.font.name = "Arial"
mr.italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1. Executive Summary")
body(doc,
    "Shot Value Over Expected (SVOE) is a shot-level NBA metric that estimates how many points "
    "a player or team should score on a given shot — based purely on information available before "
    "the shot is taken — and compares that expectation to what actually happened. "
    "A positive SVOE indicates that the player converted shots at a higher rate than expected "
    "given the shot's context; a negative SVOE indicates underperformance."
)
body(doc,
    "Unlike raw field goal percentage, SVOE adjusts for shot quality. A player who only takes "
    "corner threes and restricted-area layups is expected to score more per shot than one who "
    "takes contested mid-range jumpers. SVOE isolates shot-making ability from shot selection."
)

# ── critical limitation callout box ──────────────────────────────────────────
warn_tbl = doc.add_table(rows=1, cols=1)
warn_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
warn_cell = warn_tbl.cell(0, 0)
set_cell_bg(warn_cell, "FFF3CD")
set_cell_border(warn_cell,
    top={"val": "single", "sz": 12, "color": "E65100"},
    bottom={"val": "single", "sz": 12, "color": "E65100"},
    left={"val": "single", "sz": 12, "color": "E65100"},
    right={"val": "single", "sz": 12, "color": "E65100"},
)
warn_p = warn_cell.paragraphs[0]
warn_p.paragraph_format.space_before = Pt(10)
warn_p.paragraph_format.space_after = Pt(4)
warn_p.paragraph_format.left_indent = Inches(0.1)
warn_p.paragraph_format.right_indent = Inches(0.1)
wr1 = warn_p.add_run("IMPORTANT — What This Model Does NOT Know")
wr1.bold = True
wr1.font.size = Pt(12)
wr1.font.color.rgb = RGBColor(0xBF, 0x36, 0x00)
wr1.font.name = "Arial"

warn_p2 = warn_cell.add_paragraph()
warn_p2.paragraph_format.space_before = Pt(4)
warn_p2.paragraph_format.space_after = Pt(10)
warn_p2.paragraph_format.left_indent = Inches(0.1)
warn_p2.paragraph_format.right_indent = Inches(0.1)
wr2 = warn_p2.add_run(
    "This model does NOT know where the defender is. It cannot see whether a shot is "
    "open, lightly contested, or tightly guarded. Two shots taken from the exact same "
    "spot in the exact same game situation — one wide open, one with a hand in the face "
    "— receive identical Expected Points from this model.\n\n"
    "As a result, SVOE will naturally reward players who get open looks more often, and "
    "will understate the difficulty of shots taken against elite defenders or under heavy "
    "pressure. Any SVOE number should be read with this limitation in mind: it measures "
    "shot-making above a location- and context-based baseline, not above a fully "
    "contest-adjusted baseline."
)
wr2.font.size = Pt(11)
wr2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
wr2.font.name = "Arial"

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
#  2. CORE FORMULA
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2. Core Formula")

formula_box(doc, "SVOE  =  Actual Points  −  Expected Points")

body(doc, "Where each term is defined as:")

data_table(doc,
    ["Term", "Formula", "Description"],
    [
        ["Shot Value", "2 or 3", "2 for two-point attempts; 3 for three-point attempts"],
        ["Expected Points", "P(make) × Shot Value", "Model-predicted make probability multiplied by shot value"],
        ["Actual Points", "SHOT_MADE_FLAG × Shot Value", "1 × shot value if made; 0 × shot value if missed"],
        ["SVOE", "Actual Points − Expected Points", "Positive = over-performed; Negative = under-performed"],
    ],
    col_widths=[1.5, 2.2, 3.1],
)

body(doc,
    "SVOE is computed at the individual shot level and then aggregated across any desired "
    "grouping — player, team, season, shot zone, or action type. The most useful summary "
    "statistic is SVOE per 100 shots, which normalises for volume and allows fair comparison "
    "across players with different usage levels."
)

formula_box(doc, "SVOE / 100  =  (Total SVOE  /  Attempts)  ×  100",
            "Rate metric that adjusts for shot volume.")

# ══════════════════════════════════════════════════════════════════════════════
#  3. DATA COLLECTION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3. Data Collection")

heading(doc, "3.1  Source", level=2)
body(doc,
    "Shot data is sourced from the NBA Stats API via the open-source nba_api Python library. "
    "The ShotChartDetail endpoint returns one row per field goal attempt, including spatial "
    "coordinates, shot context, and the result."
)
bullet(doc, "Endpoint: ShotChartDetail (context_measure_simple = 'FGA')")
bullet(doc, "Seasons: 2022-23, 2023-24, 2024-25, 2025-26")
bullet(doc, "Season types: Regular Season and Playoffs")
bullet(doc, "Scope: all 30 NBA teams (queried per team to manage API rate limits)")

heading(doc, "3.2  Caching", level=2)
body(doc,
    "To avoid repeated API calls, data is cached at two levels:"
)
bullet(doc, "Per-team CSV files in data/raw/ (one file per team × season × season type)")
bullet(doc, "Merged Parquet files in data/processed/ (one file per season × season type)")
body(doc,
    "On subsequent runs, cached files are loaded directly and the API is not contacted. "
    "A --refresh-data flag exists to force a fresh download.",
    color=GREY, italic=True
)

heading(doc, "3.3  Raw Columns Used", level=2)
body(doc, "The following columns are retained from the ShotChartDetail response:")

data_table(doc,
    ["Column", "Description"],
    [
        ["GAME_ID / GAME_DATE", "Game identifier and date"],
        ["PLAYER_ID / PLAYER_NAME", "Shooter identity"],
        ["TEAM_ID / TEAM_NAME", "Shooting team"],
        ["HTM / VTM", "Home and visitor team abbreviations"],
        ["PERIOD", "Quarter or overtime period (1–4+)"],
        ["MINUTES_REMAINING / SECONDS_REMAINING", "Game clock at time of shot"],
        ["SHOT_TYPE", "'2PT Field Goal' or '3PT Field Goal'"],
        ["ACTION_TYPE", "Shot action: Jump Shot, Layup, Dunk, Hook Shot, etc."],
        ["SHOT_ZONE_BASIC / AREA / RANGE", "Zone classification at three levels of granularity"],
        ["SHOT_DISTANCE", "Distance from basket in feet"],
        ["LOC_X / LOC_Y", "Court coordinates in tenths of feet (origin = basket)"],
        ["SHOT_MADE_FLAG", "Binary outcome: 1 = made, 0 = missed"],
    ],
    col_widths=[2.8, 4.5],
)

body(doc,
    "No columns that reveal the shot result (other than SHOT_MADE_FLAG, which is the target "
    "variable) are used as model features. SHOT_MADE_FLAG is never included as a predictor.",
    italic=True, color=GREY
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4. Feature Engineering")

body(doc,
    "All features are derived solely from information available before the shot is taken. "
    "The following derived columns are added during preprocessing:"
)

data_table(doc,
    ["Derived Feature", "Source Columns", "Logic"],
    [
        ["SHOT_VALUE", "SHOT_TYPE", "3 if '3PT' in SHOT_TYPE, else 2"],
        ["TIME_REMAINING_SECS", "MINUTES_REMAINING, SECONDS_REMAINING", "Minutes × 60 + Seconds"],
        ["IS_HOME", "TEAM_ID, HTM", "1 if shooting team's abbreviation matches HTM"],
        ["OPPONENT_TEAM_ID", "HTM, VTM, IS_HOME", "Abbreviation of the opposing team mapped to team ID"],
        ["GAME_HALF", "GAME_DATE, SEASON", "H1 / H2 split at median game date within each season"],
    ],
    col_widths=[1.9, 2.3, 3.1],
)

heading(doc, "4.1  Model Feature Set", level=2)
body(doc, "Thirteen features are passed to each model:")

data_table(doc,
    ["Feature", "Type", "Rationale"],
    [
        ["LOC_X", "Numeric", "Lateral position on court (origin = basket center)"],
        ["LOC_Y", "Numeric", "Depth from basket toward mid-court"],
        ["SHOT_DISTANCE", "Numeric", "Distance in feet; highly predictive of make probability"],
        ["PERIOD", "Numeric", "Shot difficulty may vary by period / fatigue"],
        ["TIME_REMAINING_SECS", "Numeric", "Desperation shots near the buzzer have different profiles"],
        ["IS_HOME", "Binary", "Home/away splits capture crowd and travel effects"],
        ["TEAM_ID", "Categorical", "Captures team-level shooting tendencies"],
        ["OPPONENT_TEAM_ID", "Categorical", "Captures defensive quality of the defending team"],
        ["SHOT_TYPE", "Categorical", "2PT vs 3PT — strong prior on make probability"],
        ["ACTION_TYPE", "Categorical", "Layups and dunks make at much higher rates than jump shots"],
        ["SHOT_ZONE_BASIC", "Categorical", "Zone-level location (Restricted Area, Mid-Range, etc.)"],
        ["SHOT_ZONE_AREA", "Categorical", "Left/right/center area of the court"],
        ["SHOT_ZONE_RANGE", "Categorical", "Distance band within each zone"],
    ],
    col_widths=[1.8, 1.1, 4.4],
)

body(doc,
    "Note: PLAYER_ID is deliberately excluded. Including it would cause the model to absorb "
    "each player's historical make rate, reducing their SVOE toward zero. The goal is to "
    "measure performance above the contextual baseline, not to predict each player's outcome "
    "based on their identity.",
    italic=True, color=GREY
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. MODELLING
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "5. Modelling")

heading(doc, "5.1  Preprocessing Pipeline", level=2)
body(doc,
    "All feature transformation is handled by a scikit-learn ColumnTransformer that is "
    "embedded inside each model's Pipeline object:"
)
bullet(doc, "Numeric features: StandardScaler (mean=0, std=1) — harmless for trees, required for logistic regression")
bullet(doc, "Categorical features: OrdinalEncoder with handle_unknown='use_encoded_value' (unknown=-1)")
body(doc,
    "Saving each model as a single joblib file preserves the entire preprocessing chain, "
    "ensuring consistent transformation at inference time."
)

heading(doc, "5.2  Train / Calibration / Test Split", level=2)
body(doc,
    "The data is split chronologically to avoid future data leaking into the model:"
)
bullet(doc, "Training set: all seasons except the most recent (e.g., 2022-23 through 2024-25)")
bullet(doc, "Calibration set: 20% random hold-out from the training set")
bullet(doc, "Test set: the most recent season (e.g., 2025-26)")
body(doc,
    "Chronological splitting is essential: a random split would allow the model to learn "
    "from shots in the same season it is evaluated on, inflating performance metrics."
)

heading(doc, "5.3  Models Trained", level=2)

data_table(doc,
    ["Model", "Key Parameters", "Notes"],
    [
        ["Logistic Regression", "C=0.5, max_iter=1000, class_weight='balanced'", "Baseline. Naturally well-calibrated. Fast to train."],
        ["Random Forest", "300 trees, max_depth=9, min_samples_leaf=30", "Strong non-linear model. Isotonic calibration applied post-fit."],
        ["XGBoost", "500 trees, depth=5, lr=0.04, subsample=0.8", "Best performer. Gradient boosting captures complex interactions."],
    ],
    col_widths=[1.6, 2.8, 2.9],
)

heading(doc, "5.4  Probability Calibration", level=2)
body(doc,
    "Tree-based models (Random Forest and XGBoost) produce probabilities that are not always "
    "well-calibrated — they tend to push predictions toward 0 and 1. Poor calibration leads "
    "to systematically biased Expected Points, which corrupts the SVOE calculation."
)
body(doc,
    "After training each tree model on the training set, an IsotonicRegression calibrator is "
    "fit on the held-out calibration set. The calibrator maps raw predicted probabilities to "
    "empirical make rates. The final model is a two-step pipeline:"
)

formula_box(doc,
    "P_calibrated  =  IsotonicRegression( P_raw( features ) )",
    "Isotonic regression is monotone — it preserves the rank ordering of predictions while correcting the scale."
)

body(doc,
    "Logistic Regression does not require this step because its log-odds formulation already "
    "produces well-calibrated probabilities. Calibration quality is verified visually with a "
    "calibration curve comparing predicted probabilities to actual make rates across binned predictions."
)

# ══════════════════════════════════════════════════════════════════════════════
#  6. MODEL EVALUATION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "6. Model Evaluation")

body(doc, "Three metrics are computed on the held-out test set for each model:")

data_table(doc,
    ["Metric", "Formula", "Interpretation"],
    [
        ["Log Loss", "−(1/N) Σ [y log(p) + (1−y) log(1−p)]", "Overall probability quality. Lower is better. Sensitive to overconfident wrong predictions."],
        ["Brier Score", "(1/N) Σ (p − y)²", "Mean squared error of probabilities. Lower = better calibration. Primary model selection criterion."],
        ["ROC-AUC", "Area under ROC curve", "Discrimination: how well the model ranks makes above misses. 0.5 = random, 1.0 = perfect."],
    ],
    col_widths=[1.4, 2.5, 3.4],
)

body(doc,
    "The model with the lowest Brier score is selected as the best model. Brier score is "
    "preferred over log loss for this application because it penalises miscalibration in a "
    "way that directly corresponds to errors in Expected Points — a 0.05 error in predicted "
    "probability translates to a 0.05 × shot_value error in Expected Points per shot."
)

body(doc,
    "Results from the current model run (XGBoost selected as best):",
    color=GREY, italic=True
)

data_table(doc,
    ["Model", "Log Loss", "Brier Score", "ROC-AUC"],
    [
        ["Logistic Regression", "~0.66", "~0.232", "~0.638"],
        ["Random Forest", "~0.64", "~0.228", "~0.652"],
        ["XGBoost ✓", "0.637", "0.2242", "0.6604"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5],
)

body(doc,
    "An AUC of 0.66 reflects the inherent difficulty of predicting basketball shots. "
    "Shots are highly stochastic — even the best shooters make fewer than 50% of attempts "
    "from distance. The model captures the signal that exists (location, shot type, context) "
    "without overfitting to noise.",
    color=GREY, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. SVOE CALCULATION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "7. SVOE Calculation")

body(doc,
    "After selecting the best model, it is applied to every shot in the full dataset "
    "(all seasons, all teams). Three columns are added to each shot record:"
)

data_table(doc,
    ["Column", "Formula"],
    [
        ["PRED_PROB", "Best model's predict_proba output (P of make)"],
        ["EXPECTED_POINTS", "PRED_PROB × SHOT_VALUE"],
        ["ACTUAL_POINTS", "SHOT_MADE_FLAG × SHOT_VALUE"],
        ["SVOE", "ACTUAL_POINTS − EXPECTED_POINTS"],
    ],
    col_widths=[2.0, 5.3],
)

body(doc,
    "These four columns, combined with the shot metadata, form the final dataset stored in "
    "data/processed/shots_with_svoe.parquet. The dashboard reads this file at startup."
)

heading(doc, "7.1  Aggregation", level=2)
body(doc, "SVOE can be aggregated over any grouping. The standard summary statistics are:")

data_table(doc,
    ["Statistic", "Formula", "Meaning"],
    [
        ["Total SVOE", "Σ SVOE", "Cumulative over/under performance in points"],
        ["SVOE / 100", "(Σ SVOE / N) × 100", "Rate per 100 shots — adjusts for volume"],
        ["EP / Shot", "Σ Expected Points / N", "Shot quality: how good are the looks being generated?"],
        ["AP / Shot", "Σ Actual Points / N", "Shooting efficiency: how many points per attempt?"],
    ],
    col_widths=[1.5, 2.3, 3.5],
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. KEY FINDINGS
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "8. Key Findings by Shot Zone")

body(doc,
    "The following zone-level results are observed across all four seasons "
    "(2022-23 through 2025-26):"
)

data_table(doc,
    ["Shot Zone", "FGA", "EP/Shot", "AP/Shot", "SVOE/100", "Interpretation"],
    [
        ["Above the Break 3", "278,941", "1.058", "1.062", "+0.31", "Neutral — model well-calibrated on high-volume zone"],
        ["Restricted Area", "268,605", "1.339", "1.330", "−0.90", "Model over-estimates rim makes (can't see contest)"],
        ["In The Paint (Non-RA)", "185,686", "0.887", "0.886", "−0.11", "Slightly below expectation — contested paint shots"],
        ["Mid-Range", "100,844", "0.839", "0.838", "−0.14", "Marginally below expectation — tough pull-ups"],
        ["Left Corner 3", "50,104", "1.140", "1.155", "+1.52", "Model underestimates corner 3 — shorter distance"],
        ["Right Corner 3", "46,233", "1.152", "1.163", "+1.03", "Same effect, right side"],
    ],
    col_widths=[1.7, 0.85, 0.75, 0.75, 0.75, 2.5],
)

body(doc, "These findings align with established basketball analytics intuition:")
bullet(doc,
    "Corner 3s outperform expectation (+1.52, +1.03) because they are shorter "
    "than above-the-break 3s (~22 ft vs ~23.75 ft) and are frequently left open "
    "by defences rotating to protect the paint."
)
bullet(doc,
    "The Restricted Area underperforms (−0.90) because the model assigns high make "
    "probability to all near-rim attempts without knowing whether the shot is contested. "
    "Blocked shots and heavily contested layups drag actual make rates below predictions."
)
bullet(doc,
    "Mid-range shots are slightly below expectation (−0.14), consistent with the "
    "widely documented difficulty of pull-up mid-range jumpers in the modern NBA."
)
bullet(doc,
    "The league-wide SVOE is approximately zero (expected by construction, since the "
    "model is trained on the same population it scores)."
)

# ══════════════════════════════════════════════════════════════════════════════
#  9. LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "9. Limitations & Caveats")

bullet(doc,
    "No shot contest data: the model cannot observe whether a shot is open or "
    "tightly contested. This is the primary source of model error, particularly "
    "at the rim and on pull-up jumpers."
)
bullet(doc,
    "No defender identity: the quality of the closest defender is not captured. "
    "A layup over a 7-foot shot-blocker and an unguarded breakaway are treated identically."
)
bullet(doc,
    "No fatigue or injury context: back-to-back games, minutes load, and player "
    "health status are not modelled."
)
bullet(doc,
    "Small-sample noise: SVOE/100 is volatile for players with fewer than ~200 "
    "attempts. The dashboard requires a minimum attempts filter (default 100) to "
    "reduce noise in the leaderboards."
)
bullet(doc,
    "Model stationarity: the XGBoost model is trained on historical data. Rule "
    "changes, player development, or shifts in league-wide shot selection may reduce "
    "accuracy over time. The model should be retrained annually."
)
bullet(doc,
    "Sustainability: high SVOE/100 in a short window (e.g., the first half of a "
    "season) has meaningful but limited predictive power for the second half. "
    "The Sustainability tab in the dashboard quantifies year-over-year stability."
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. TECHNICAL STACK
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "10. Technical Stack")

data_table(doc,
    ["Component", "Library / Tool", "Version"],
    [
        ["Data retrieval", "nba_api", "≥ 1.4.1"],
        ["Data processing", "pandas, numpy", "≥ 2.0, ≥ 1.24"],
        ["Machine learning", "scikit-learn", "≥ 1.3"],
        ["Gradient boosting", "XGBoost", "≥ 2.0"],
        ["Data storage", "Apache Parquet (pyarrow)", "≥ 13.0"],
        ["Visualisation", "Plotly", "≥ 5.17"],
        ["Dashboard", "Streamlit", "≥ 1.28"],
        ["Model persistence", "joblib", "≥ 1.3"],
    ],
    col_widths=[2.0, 2.5, 2.8],
)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
add_horizontal_rule(doc, color="CCCCCC", thickness=6)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    "SVOE Analytics  •  Data via NBA Stats API (nba_api)  •  "
    "Seasons 2022-23 through 2025-26"
)
fr.font.size = Pt(9)
fr.font.color.rgb = GREY
fr.font.name = "Arial"
fr.italic = True

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT.resolve()}")
